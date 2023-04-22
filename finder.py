import re
import requests
from docx import Document
from tkinter import *
from tkinter import filedialog
from bs4 import BeautifulSoup


class App:
    def __init__(self, master):
        self.master = master
        master.title("Web Search Tool")

        self.search_terms_label = Label(master, text="Enter search term(s):")
        self.search_terms_label.pack()

        self.search_terms_entry = Entry(master)
        self.search_terms_entry.pack()

        self.url_label = Label(master, text="Enter website URL(s), separated by commas:")
        self.url_label.pack()

        self.url_entry = Entry(master)
        self.url_entry.pack()

        self.run_button = Button(master, text="Run Search", command=self.run_search)
        self.run_button.pack()

        self.quit_button = Button(master, text="Quit", command=master.quit)
        self.quit_button.pack()

    def run_search(self):
        search_terms = self.search_terms_entry.get().split(',')
        urls = self.url_entry.get().split(',')
        document = Document()

        for url in urls:
            url = url.strip()
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
            except requests.exceptions.Timeout as e:
                print(f"Request timed out for {url}: {e}")
                continue
            except requests.exceptions.RequestException as e:
                print(f"Failed to retrieve content from {url}: {e}")
                continue

            content_type = response.headers.get('Content-Type')
            if 'text' in content_type or 'html' in content_type:
                content = response.text
                soup = BeautifulSoup(content, 'html.parser')
                paragraphs = soup.find_all('p')
                for paragraph in paragraphs:
                    text = paragraph.get_text()
                    for search_term in search_terms:
                        # Use a regular expression to find the search term, ignoring case and allowing for variations
                        pattern = re.compile(search_term, re.IGNORECASE)
                        if pattern.search(text):
                            document.add_paragraph(text)
                            break  # If a match is found, don't continue searching for other search terms
                print(f"Found search results in {url}.")
            else:
                print(f"Content type not supported for {url}: {content_type}")
                continue

        if len(document.paragraphs) > 0:
            filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Microsoft Word Document", "*.docx")])
            document.save(filename)
            print(f"Saved search results to {filename}.")
        else:
            print("No search results found.")

root = Tk()
app = App(root)
root.mainloop()
